#!/usr/bin/env python3
"""
Сборка DOCUMENTATION.docx и HOW_IT_WORKS.docx из markdown в корне репозитория.

  pip install python-docx
  python scripts/build_docs_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
DOCS_OUT = REPO_ROOT / "docs"

SOURCES = (
    ("DOCUMENTATION.md", "DOCUMENTATION.docx", "Афиша Sirius + помощник — руководство"),
    ("HOW_IT_WORKS.md", "HOW_IT_WORKS.docx", "Афиша Sirius + помощник — как устроен проект"),
)


def _ensure_docx():
    try:
        from docx import Document  # noqa: F401
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_LINE_SPACING
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("Установите зависимость: pip install python-docx", file=sys.stderr)
        sys.exit(1)
    return Document, Cm, Pt, RGBColor, WD_LINE_SPACING, WD_STYLE_TYPE


def _configure_styles(doc, Cm, Pt, WD_LINE_SPACING, WD_STYLE_TYPE):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for level, size in ((1, 18), (2, 14), (3, 12)):
        name = f"Heading {level}"
        h = doc.styles[name]
        h.font.name = "Calibri"
        h.font.bold = True
        h.font.size = Pt(size)
        h.font.color.rgb = None
        h.paragraph_format.space_before = Pt(12 if level > 1 else 0)
        h.paragraph_format.space_after = Pt(6)

    if "CodeBlock" not in [s.name for s in doc.styles]:
        try:
            code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            code_style = doc.styles["CodeBlock"]
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Cm(0.5)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)


def _add_runs(paragraph, text: str, *, code_font: bool = False):
    from docx.shared import Pt

    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            if code_font:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
        chunk = m.group(0)
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif chunk.startswith("*") and chunk.endswith("*"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if code_font:
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def _parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", line))


def markdown_to_docx(md_path: Path, docx_path: Path, subtitle: str) -> None:
    Document, Cm, Pt, RGBColor, WD_LINE_SPACING, WD_STYLE_TYPE = _ensure_docx()
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    _configure_styles(doc, Cm, Pt, WD_LINE_SPACING, WD_STYLE_TYPE)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(subtitle)
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"Источник: {md_path.name}")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    lines = text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    skip_first_h1 = True

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="CodeBlock")
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph("—" * 40)
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("## "):
            if skip_first_h1:
                skip_first_h1 = False
                i += 1
                continue
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and _is_table_sep(lines[i + 1]):
            headers = _parse_table_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_parse_table_row(lines[i]))
                i += 1
            ncol = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=ncol)
            table.style = "Table Grid"
            for c, h in enumerate(headers):
                cell = table.rows[0].cells[c]
                cell.text = ""
                run = cell.paragraphs[0].add_run(h)
                run.bold = True
            for r_idx, row in enumerate(rows):
                for c in range(ncol):
                    val = row[c] if c < len(row) else ""
                    table.rows[r_idx + 1].cells[c].text = val
            doc.add_paragraph()
            continue

        if line.strip().startswith("- ") or line.strip().startswith("* "):
            item = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, item)
            i += 1
            continue

        m_num = re.match(r"^(\d+)\.\s+(.*)$", line.strip())
        if m_num:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, m_num.group(2))
            i += 1
            continue

        if line.strip().startswith("```mermaid"):
            note = doc.add_paragraph()
            _add_runs(
                note,
                "[Диаграмма Mermaid — см. электронный файл HOW_IT_WORKS.md или репозиторий.]",
            )
            note.paragraph_format.left_indent = Cm(0.5)
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            if i < len(lines):
                i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        _add_runs(p, line.strip())
        i += 1

    if in_code and code_lines:
        p = doc.add_paragraph(style="CodeBlock")
        p.add_run("\n".join(code_lines)).font.name = "Consolas"

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"OK  {docx_path.relative_to(REPO_ROOT)}")


def main() -> None:
    for md_name, docx_name, subtitle in SOURCES:
        md_path = REPO_ROOT / md_name
        if not md_path.is_file():
            print(f"Нет файла: {md_path}", file=sys.stderr)
            sys.exit(1)
        markdown_to_docx(md_path, DOCS_OUT / docx_name, subtitle)


if __name__ == "__main__":
    main()
